#!/usr/bin/env python3
"""Génère la facture Word pour la Commune (OCS / Œuvre des Colonies Scolaires)
à partir d'un CSV d'inscriptions.

CSV attendu (en-tête obligatoire), une ligne par inscription à FACTURER :
    semaine,nom,prenom,formule,montant
- semaine : libellé de regroupement, p.ex. "Semaine 2" (les lignes sont
  regroupées et affichées dans l'ordre d'apparition des semaines).
- montant : entier en euros pour cette inscription (120 en général,
  96 pour une semaine écourtée par un jour férié). Voir SKILL.md.

Ne mettre dans le CSV QUE les inscriptions payées par la commune
(les enfants marqués "(COMMUNE)"). Le filtrage se fait en amont.

Usage :
    python build_facture.py inscriptions.csv \
        --numero OCS027 \
        --concerne "stages découverte du tennis 3-5 ans, mini-tennis 5-7 ans et tennis/multisports, Toussaint 2026 (25/10/2026-31/10/2026)" \
        --sortie Facture_OCS027.docx
"""
import argparse
import csv
import os
from collections import OrderedDict

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from euros_en_lettres import euros_en_lettres

ASSETS = os.path.join(os.path.dirname(__file__), "..", "assets")
LOGO = os.path.join(ASSETS, "entete_lambermont.png")

# --- Coordonnées fixes (RTC Lambermont -> Commune de Schaerbeek) -------------
DESTINATAIRE = ("Commune de Schaerbeek ASBL OCS, Œuvre des Colonies Scolaire "
                "78-80, 1030 Schaerbeek")
IBAN = "BE48068942393827"
SIGNATAIRE = "Thibault Duchène"
FONCTION = "Directeur de l'école de tennis du Lambermont"
# Ordre d'affichage des formules à l'intérieur d'une semaine
ORDRE_FORMULES = [
    "Stage Découverte du Tennis 3 - 5ans",
    "Stage Mini-Tennis - 5 - 7ans",
    "Stage Tennis - Multisports - 8 à 18 ans",
    "Stage Compétition - 7 à 18 ans",
]


def _ordre(formule: str) -> int:
    for i, f in enumerate(ORDRE_FORMULES):
        if formule.strip() == f:
            return i
    return len(ORDRE_FORMULES)  # formules inconnues en fin de liste


def lire_csv(path):
    semaines = OrderedDict()
    with open(path, encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            sem = row["semaine"].strip()
            semaines.setdefault(sem, []).append({
                "nom": row["nom"].strip(),
                "prenom": row["prenom"].strip(),
                "formule": row["formule"].strip(),
                "montant": int(str(row["montant"]).strip()),
            })
    # tri interne : par formule (ordre métier) puis par nom
    for lignes in semaines.values():
        lignes.sort(key=lambda r: (_ordre(r["formule"]), r["nom"], r["prenom"]))
    return semaines


def _p(doc, texte="", *, gras=False, taille=11, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if texte:
        run = p.add_run(texte)
        run.bold = gras
        run.font.size = Pt(taille)
    return p


def _set_cell(cell, texte, *, gras=False, taille=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(texte))
    run.bold = gras
    run.font.size = Pt(taille)


def construire(semaines, numero, concerne, sortie):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    if os.path.exists(LOGO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(LOGO, width=Cm(16))

    _p(doc, f"Facture n° {numero},", gras=True, taille=13)
    _p(doc, "A l'attention de:", gras=True)
    _p(doc, DESTINATAIRE)
    _p(doc, f"Concerne: Facture du RTC Lambermont, {concerne}")
    _p(doc)
    _p(doc, "Chère Madame, Cher Monsieur,")

    total = sum(l["montant"] for lignes in semaines.values() for l in lignes)
    _p(doc, f"Merci de verser la somme de {total} euros sur le compte {IBAN} "
            "correspondant aux inscriptions suivantes:")

    for sem, lignes in semaines.items():
        _p(doc, f"{sem}:", gras=True, space_after=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for c, titre in zip(table.rows[0].cells, ["Nom", "Prénom", "Formule", "Solde dû"]):
            _set_cell(c, titre, gras=True)
        for l in lignes:
            cells = table.add_row().cells
            _set_cell(cells[0], l["nom"])
            _set_cell(cells[1], l["prenom"])
            _set_cell(cells[2], l["formule"])
            _set_cell(cells[3], l["montant"])
        sous_total = sum(l["montant"] for l in lignes)
        st = doc.add_paragraph()
        st.paragraph_format.space_before = Pt(2)
        r1 = st.add_run("Sous-total  ")
        r1.bold = True
        r2 = st.add_run(str(sous_total))
        r2.bold = True
        doc.add_paragraph()

    n = len(semaines)
    mot_sem = "semaine" if n == 1 else "semaines"
    _p(doc, f"Total sur les {n} {mot_sem}: {total} euros", gras=True)
    _p(doc, f"Certifié sincère et véritable, arrêté la somme de {euros_en_lettres(total)}")
    _p(doc)
    _p(doc, "Vous remerciant d'avance,")
    _p(doc, SIGNATAIRE, space_after=0)
    _p(doc, FONCTION)

    doc.save(sortie)
    return total, n


def main():
    ap = argparse.ArgumentParser(description="Génère la facture Word pour la Commune")
    ap.add_argument("csv", help="CSV des inscriptions à facturer")
    ap.add_argument("--numero", required=True, help="Numéro de facture, p.ex. OCS027")
    ap.add_argument("--concerne", required=True,
                    help="Description après 'Facture du RTC Lambermont, '")
    ap.add_argument("--sortie", required=True, help="Chemin du .docx de sortie")
    args = ap.parse_args()

    semaines = lire_csv(args.csv)
    total, n = construire(semaines, args.numero, args.concerne, args.sortie)
    print(f"OK  {args.sortie}")
    print(f"    {sum(len(v) for v in semaines.values())} inscriptions sur {n} semaine(s)")
    print(f"    Total: {total} euros ({euros_en_lettres(total)})")


if __name__ == "__main__":
    main()
